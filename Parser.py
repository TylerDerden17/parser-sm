import os
import re
import requests
from PIL import Image
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn  # Явное указание шрифта


def normalize_url(input_url):
    match = re.search(r'/(\d+)', input_url)
    if match:
        article_id = match.group(1)
        return f"https://smart-lab.ru/mobile/topic/{article_id}"
    else:
        raise ValueError("Не удалось извлечь ID статьи из ссылки.")


def download_image(img_url, img_path):
    try:
        img_data = requests.get(img_url)
        img_data.raise_for_status()
        with open(img_path, 'wb') as handler:
            handler.write(img_data.content)
        return img_path
    except requests.exceptions.RequestException as e:
        print(f"❌ Ошибка при скачивании изображения {img_url}: {e}")
        return None


def convert_image_to_supported_format(img_path):
    if img_path.lower().endswith(".webp"):
        try:
            img = Image.open(img_path)
            img_path_new = img_path.replace(".webp", ".png")
            img.save(img_path_new, "PNG")
            os.remove(img_path)
            print(f"🔁 Конвертировано: {img_path_new}")
            return img_path_new
        except Exception as e:
            print(f"❌ Ошибка конвертации {img_path}: {e}")
    return img_path


def save_article_to_docx(raw_url):
    url = normalize_url(raw_url)

    try:
        response = requests.get(url)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(f"❌ Ошибка загрузки статьи: {e}")
        return

    soup = BeautifulSoup(response.text, "html.parser")

    title_tag = soup.select_one("div.post-card__title-wrap")
    title = title_tag.get_text(strip=True) if title_tag else "Статья"
    safe_title = re.sub(r'[\\/*?:"<>|]', "", title).replace('.', '_').replace(',', '_').replace(' ', '_')

    article_dir = os.path.join("articles", safe_title)
    os.makedirs(article_dir, exist_ok=True)

    images_dir = os.path.join(article_dir, "images")
    os.makedirs(images_dir, exist_ok=True)

    doc = Document()

    # 🔷 Заголовок
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.font.name = 'Arial Black'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Black')
    run.font.size = Pt(18)

    content = soup.select_one("div.post-card__text")
    if not content:
        print("⚠️ Контент статьи не найден.")
        return

    # 🔹 Добавление контента с учетом всех элементов (p, h1, h2, ul, ol, li, img и т.д.)
    for element in content.find_all(True):  # True означает обработку всех тегов
        if element.name == "p":  # Текстовый параграф
            text = element.get_text(strip=True)
            if text:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(text)
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                run.font.size = Pt(12)

        elif element.name == "h1":  # Заголовок 1 уровня
            text = element.get_text(strip=True)
            if text:
                heading = doc.add_paragraph()
                run = heading.add_run(text)
                run.font.name = 'Arial Black'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Black')
                run.font.size = Pt(16)

        elif element.name == "h2":  # Заголовок 2 уровня
            text = element.get_text(strip=True)
            if text:
                heading = doc.add_paragraph()
                run = heading.add_run(text)
                run.font.name = 'Arial Black'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Black')
                run.font.size = Pt(14)

        elif element.name == "h3":  # Заголовок 3 уровня
            text = element.get_text(strip=True)
            if text:
                heading = doc.add_paragraph()
                run = heading.add_run(text)
                run.font.name = 'Arial Black'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Black')
                run.font.size = Pt(12)

        elif element.name == "ul" or element.name == "ol":  # Список
            for li in element.find_all("li"):
                list_item = li.get_text(strip=True)
                if list_item:
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(f"• {list_item}")  # Список с маркером

        elif element.name == "img":  # Изображение
            img_url = element.get("src")
            if not img_url:
                continue

            if not img_url.startswith("http"):
                img_url = "https://smart-lab.ru" + img_url

            img_name = img_url.split("/")[-1].split("?")[0]
            img_path = os.path.join(images_dir, img_name)

            downloaded_img_path = download_image(img_url, img_path)
            if downloaded_img_path:
                img_path = convert_image_to_supported_format(downloaded_img_path)
                try:
                    doc.add_picture(img_path, width=Inches(5.0))
                    print(f"✅ Вставлено изображение: {img_path}")
                except Exception as e:
                    print(f"❌ Ошибка при вставке изображения {img_path}: {e}")

    docx_path = os.path.join(article_dir, f"{safe_title}.docx")
    doc.save(docx_path)
    print(f"\n📄 Статья успешно сохранена: {docx_path}")
